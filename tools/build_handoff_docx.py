from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "组员交接说明.md"
OUTPUT = ROOT / "deliverables" / "CUMCM2025A_组员交接说明_20260826.docx"

INK = "20364B"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "64748B"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
BLACK = "000000"
WHITE = "FFFFFF"

PAGE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run, *, size=11, bold=False, italic=False, color=BLACK, mono=False):
    western = "Consolas" if mono else "Calibri"
    eastern = "Microsoft YaHei"
    run.font.name = western
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), western)
    rpr.rFonts.set(qn("w:hAnsi"), western)
    rpr.rFonts.set(qn("w:eastAsia"), eastern)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    if sum(widths) != PAGE_WIDTH_DXA:
        raise ValueError(f"table widths must sum to {PAGE_WIDTH_DXA}: {widths}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(PAGE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_numbering(doc, *, bullet=False):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if bullet else "%1.")
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    ppr.extend([tabs, ind])
    lvl.extend([start, num_fmt, lvl_text, suff, ppr])
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    ref = OxmlElement("w:abstractNumId")
    ref.set(qn("w:val"), str(abstract_id))
    num.append(ref)
    numbering.append(num)
    return num_id


def set_num(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    ppr.append(num_pr)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)
    run = paragraph.add_run(" 页")
    set_run_font(run, size=9, color=MUTED)


def configure(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    hp = section.header.paragraphs[0]
    hp.paragraph_format.space_after = Pt(0)
    run = hp.add_run("CUMCM 2025 A｜组员交接说明")
    set_run_font(run, size=9, color=MUTED)
    add_page_number(section.footer.paragraphs[0])


def add_inline(paragraph, text, *, size=10.6, color=BLACK):
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|<https?://[^>]+>)")
    for part in pattern.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=size, bold=True, color=color)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=size - 0.4, color=DARK_BLUE, mono=True)
        else:
            cleaned = part[1:-1] if part.startswith("<http") and part.endswith(">") else part
            run = paragraph.add_run(cleaned)
            set_run_font(run, size=size, color=color)


def add_body_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    add_inline(p, text)


def add_list_item(doc, text, num_id):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    set_num(p, num_id)
    add_inline(p, text)


def add_table(doc, rows):
    cols = len(rows[0])
    widths = {
        2: [2700, 6660],
        3: [1900, 3430, 4030],
        4: [1500, 2500, 2500, 2860],
    }.get(cols)
    if widths is None:
        base = PAGE_WIDTH_DXA // cols
        widths = [base] * cols
        widths[-1] += PAGE_WIDTH_DXA - sum(widths)

    table = doc.add_table(rows=len(rows), cols=cols, style="Table Grid")
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            cell = table.cell(i, j)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            add_inline(p, value, size=9.0, color=INK if i == 0 else BLACK)
            if i == 0:
                for run in p.runs:
                    run.bold = True
                set_cell_shading(cell, LIGHT_BLUE)
    set_table_geometry(table, widths)
    header_pr = table.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    header_pr.append(repeat)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)


def parse_markdown(doc, text):
    lines = text.splitlines()
    index = 0
    first_heading = True
    current_list_kind = None
    current_list_id = None

    while index < len(lines):
        line = lines[index].rstrip()
        if not line:
            current_list_kind = None
            current_list_id = None
            index += 1
            continue

        if line.startswith("|"):
            rows = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                cells = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
                    rows.append(cells)
                index += 1
            if rows:
                add_table(doc, rows)
            current_list_kind = None
            current_list_id = None
            continue

        heading = re.match(r"^(#{1,3})\s+(.*)$", line)
        if heading:
            level = len(heading.group(1))
            title = heading.group(2)
            if first_heading and level == 1:
                first_heading = False
                index += 1
                continue
            doc.add_heading(title, level=min(level, 3))
            current_list_kind = None
            current_list_id = None
            index += 1
            continue

        if line.startswith("- "):
            if current_list_kind != "bullet":
                current_list_kind = "bullet"
                current_list_id = add_numbering(doc, bullet=True)
            add_list_item(doc, line[2:], current_list_id)
            index += 1
            continue

        numbered = re.match(r"^\d+\.\s+(.*)$", line)
        if numbered:
            if current_list_kind != "number":
                current_list_kind = "number"
                current_list_id = add_numbering(doc, bullet=False)
            add_list_item(doc, numbered.group(1), current_list_id)
            index += 1
            continue

        add_body_paragraph(doc, line.replace("  ", ""))
        current_list_kind = None
        current_list_id = None
        index += 1


def build():
    doc = Document()
    configure(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("CUMCM 2025 A题项目交接与下一步分工")
    set_run_font(run, size=24, bold=True, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run("安泰队友负责 coding｜数院队友负责建模｜你负责 Git 整合与最终把关")
    set_run_font(run, size=12.3, color=MUTED)

    callout = doc.add_table(rows=1, cols=1)
    set_table_geometry(callout, [PAGE_WIDTH_DXA])
    set_cell_shading(callout.cell(0, 0), LIGHT_BLUE)
    p = callout.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    add_inline(
        p,
        "交接原则：不从零重做，不更换观测口径。coding 线保证可复现和结果一致，建模线保证定义、证明与结论边界，最终通过 Git 审查汇合。",
        size=10.5,
        color=INK,
    )
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)

    parse_markdown(doc, SOURCE.read_text(encoding="utf-8"))

    core = doc.core_properties
    core.title = "CUMCM 2025 A题项目交接与下一步分工"
    core.subject = "三人组 coding、建模与 Git 协作说明"
    core.author = "CUMCM 三人组"
    core.keywords = "CUMCM, 数学建模, coding, Git, 交接"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()

